from fastapi import APIRouter, Depends, Request
from fastapi.responses import JSONResponse
from docx import Document
import os
from typing import List
from pydantic import BaseModel
import json
import base64
from io import BytesIO
from jwt_manager import JWTBearer

archivo_router = APIRouter()

class Producto(BaseModel):
    cantidad: int
    producto: str
    descripcion: str
    unidad: str
    valor_unitario: str
    valor_total: str
    
class Cotizacion(BaseModel):
    fecha: str
    cliente: str
    usuario: str
    cotizacion: List[Producto]
    valor_total: str
    representante: str
    url_chatbot: str
    cargo_representante: str
    correo_asistenteia: str
    telefono: str
    nombre_plantilla: str


@archivo_router.post("/convert/", dependencies=[Depends(JWTBearer())])
async def convert_word(
    request: Request,
    cotizacion: Cotizacion
):
    # Parsear la cotización como una lista de productos
    productos = cotizacion.cotizacion

    # Directorio temporal y archivo fijo
    temp_dir = "static/temp_files"
    os.makedirs(temp_dir, exist_ok=True)
    input_path = os.path.join(temp_dir, f"{cotizacion.nombre_plantilla}")

    # Verificar si el archivo fijo existe
    if not os.path.exists(input_path):
        raise RuntimeError(f"El archivo fijo no existe: {input_path}")

    # Cargar el documento existente
    document = Document(input_path)
    # Reemplazar texto en párrafos
    for paragraph in document.paragraphs:
        if "FECHA" in paragraph.text:
            paragraph.text = paragraph.text.replace("FECHA", cotizacion.fecha)
        if "USUARIO" in paragraph.text:
            paragraph.text = paragraph.text.replace("USUARIO", cotizacion.usuario)
        if "VALORTOTAL" in paragraph.text:
            # Convertir valor_total a un número antes de formatear
            valor_total = float(cotizacion.valor_total)
            paragraph.text = paragraph.text.replace("VALORTOTAL", f"{valor_total:,.0f}")
        if "CLIENTE" in paragraph.text:
            paragraph.text = paragraph.text.replace("CLIENTE", cotizacion.cliente)
        if "REPRESENTANTE" in paragraph.text:
            paragraph.text = paragraph.text.replace("REPRESENTANTE", cotizacion.representante)
        if "URL_CHATBOT" in paragraph.text:
            paragraph.text = paragraph.text.replace("URL_CHATBOT", cotizacion.url_chatbot)
        if "CARGOREPRE" in paragraph.text:
            paragraph.text = paragraph.text.replace("CARGOREPRE", cotizacion.cargo_representante)
        if "CORREO_ASISTENTEIA" in paragraph.text:
            paragraph.text = paragraph.text.replace("CORREO_ASISTENTEIA", cotizacion.correo_asistenteia)
        if "TELEFONO" in paragraph.text:
            paragraph.text = paragraph.text.replace("TELEFONO", cotizacion.telefono)


    # Buscar la tabla de cotización en el documento
    for table in document.tables:
        if "Cantidad" in table.cell(0, 0).text:
            # Borrar filas existentes excepto la cabecera
            while len(table.rows) > 1:
                table._element.remove(table.rows[-1]._element)

            # Agregar filas con datos de productos
            for producto in productos:
                row = table.add_row()
                row.cells[0].text = str(producto.cantidad)
                row.cells[1].text = producto.producto
                row.cells[2].text = producto.descripcion
                row.cells[3].text = producto.unidad
                
                # Convertir valor_unitario y valor_total a números antes de formatear
                valor_unitario = float(producto.valor_unitario)
                valor_total = float(producto.valor_total)
                
                row.cells[4].text = f"${valor_unitario:,.0f}"
                row.cells[5].text = f"${valor_total:,.0f}"


    # Guardar el documento modificado en un buffer de bytes
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)

    # Codificar el archivo en base64
    base64_content = base64.b64encode(buffer.read()).decode('utf-8')

    # Devolver el contenido base64, el nombre y el tipo de archivo
    response = {
        "content": base64_content,
        "name": f"cotizacion_{cotizacion.cliente}_{cotizacion.usuario}.docx",
        "type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    }

    return JSONResponse(content=response)

